#!/usr/bin/env python3
"""Parse the 4 new MCQ/docx sources into structured JSON for the Q&A tab + flash deck.
Sources are QUESTION SOURCES only — answers are re-verified from official books later.
"""
import docx, re, json, sys

OUT = "/data/prometric/work/parsed_new_mcqs.json"

def clean(s):
    s = re.sub(r"[ \t]+", " ", s).strip()
    return s

def parse_july2026(path):
    d = docx.Document(path)
    paras = [clean(p.text) for p in d.paragraphs if clean(p.text)]
    items, cur = [], None
    def is_question(p):
        if len(p) < 15: return False
        if p.endswith("?") or p.endswith("? ("): return True
        if re.match(r"^(Which|What|How|When|Why|Where|A patient|During|In a|In the|The |A |An |Regarding|A composite|The image|Based on|A newborn|A panoramic|A pregnant|A tooth)", p): return True
        return False
    def is_meta(p):
        if not p: return True
        if p.startswith(("July", "ابطال", "🟡", "🟢", "🔁", "اي شخص", "يراسلني", "الاجوبة", "اتمنى", "الاسئلة", "الملف", "الملفات", "الافكار", "ممكن", "هذي", "اللي", "الصور", "الصورة")): return True
        if re.fullmatch(r"[✅✔🔁🟡🟢]+", p): return True
        return False
    for p in paras:
        if is_meta(p): continue
        if is_question(p):
            if cur: items.append(cur)
            cur = {"source": "july2026", "stem": p, "options": [], "answerIdx": None, "image": ("image" in p.lower() or "shown" in p.lower())}
        elif cur is not None:
            # plain option line, no letter prefix in this file
            if p and not is_meta(p):
                txt = p.replace("✅", "").replace("✔", "").strip()
                is_ans = "✅" in p or "✔" in p
                cur["options"].append(txt)
                if is_ans: cur["answerIdx"] = len(cur["options"]) - 1
    if cur: items.append(cur)
    return [i for i in items if len(i["options"]) >= 2]

def parse_mcq_solved(path):
    d = docx.Document(path)
    paras = [clean(p.text) for p in d.paragraphs if clean(p.text)]
    items, cur = [], None
    for p in paras:
        m = re.match(r"^(\d+)\)\s*(.+)$", p)
        if m:
            if cur: items.append(cur)
            cur = {"source": "mcq_solved", "stem": m.group(2), "options": [], "answerIdx": None, "why": "", "reference": ""}
        elif cur is not None:
            mo = re.match(r"^([A-E])\)\s*(.+?)\s*(✔ CORRECT|✔|CORRECT)?\s*$", p)
            if mo:
                txt = mo.group(2).strip()
                cur["options"].append(txt)
                if mo.group(3): cur["answerIdx"] = len(cur["options"]) - 1
            elif p.startswith("Answer:"):
                cur["answerLetter"] = p.split(":")[-1].strip()
            elif p.startswith("Why:"):
                cur["why"] = p[4:].strip()
    if cur: items.append(cur)
    # resolve answerIdx from answerLetter
    for it in items:
        if it.get("answerIdx") is None and it.get("answerLetter"):
            al = it["answerLetter"].upper()
            for i, o in enumerate(it["options"]):
                if o == it["options"][0]: pass
            # find by letter
            idx = ord(al) - ord("A")
            if 0 <= idx < len(it["options"]): it["answerIdx"] = idx
    return [i for i in items if len(i["options"]) >= 2]

def parse_bank160(path):
    d = docx.Document(path)
    paras = [clean(p.text) for p in d.paragraphs if clean(p.text)]
    items, cur = [], None
    for p in paras:
        m = re.match(r"^Q(\d+)\.\s*(.+)$", p)
        if m:
            if cur: items.append(cur)
            cur = {"source": "bank160", "qnum": int(m.group(1)), "stem": m.group(2), "options": [], "answerIdx": None, "why": "", "reference": ""}
        elif cur is not None:
            mo = re.match(r"^([A-E])\.\s*(.+?)\s*(✅|✔)?\s*$", p)
            if mo:
                txt = mo.group(2).strip()
                cur["options"].append(txt)
                if mo.group(3): cur["answerIdx"] = len(cur["options"]) - 1
            elif p.startswith("Bec:"):
                cur["why"] = p[4:].strip()
            elif p.startswith("Reference:"):
                cur["reference"] = p[10:].strip()
    if cur: items.append(cur)
    return [i for i in items if len(i["options"]) >= 2]

def parse_qa(path):
    d = docx.Document(path)
    paras = [clean(p.text) for p in d.paragraphs if clean(p.text)]
    items, cur = [], None
    for p in paras:
        m = re.match(r"^Q(\d+)\.\s*(.+)$", p)
        if m:
            if cur: items.append(cur)
            cur = {"source": "qa_answered", "qnum": int(m.group(1)), "stem": m.group(2), "options": [], "answerIdx": None, "why": "", "reference": ""}
        elif cur is not None:
            if p.startswith("Answer:"):
                cur["answer"] = p[7:].strip()
            elif p.startswith("Reference:"):
                cur["reference"] = p[10:].strip()
            elif p.startswith("Why:"):
                cur["why"] = p[4:].strip()
    if cur: items.append(cur)
    return items  # Q&A, may have no options

if __name__ == "__main__":
    src = {
        "july2026": "/data/prometric/July 2026 Questions  أبطال الدجيتال .docx",
        "mcq_solved": "/home/kalde/Downloads/WhatsApp Unknown 2026-08-05 at 04.48.57/MCQs_Solved_with_Explanations.docx",
        "bank160": "/home/kalde/Downloads/WhatsApp Unknown 2026-08-05 at 04.48.57/SDLE_BANK_ANSWERED_2026_EN.docx",
        "qa_answered": "/home/kalde/Downloads/WhatsApp Unknown 2026-08-05 at 04.48.57/SDLE_QA_Answered.docx",
    }
    out = {}
    for name, path in src.items():
        fn = {"july2026": parse_july2026, "mcq_solved": parse_mcq_solved, "bank160": parse_bank160, "qa_answered": parse_qa}[name]
        try:
            items = fn(path)
        except Exception as e:
            print(f"{name}: PARSE ERROR {e}"); items = []
        out[name] = items
        n_ok = sum(1 for i in items if i.get("answerIdx") is not None and len(i.get("options", [])) >= 2)
        print(f"{name}: {len(items)} questions | {n_ok} with options+answer")
    json.dump(out, open(OUT, "w", encoding="utf-8"), ensure_ascii=False, indent=1)
    print("saved", OUT)
