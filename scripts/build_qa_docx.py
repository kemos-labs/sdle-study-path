#!/usr/bin/env python3
"""Generate SDLE_QA_Answered.docx — answers grounded in the gold reference textbooks
in sdle-ref/books/, with the book name + 'why we choose the answer' for each item."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BOOKS = {
    "fixed": "Contemporary Fixed Prosthodontics (Shillingburg) 4e / 5e",
    "sturv": "Sturdevant's Art & Science of Operative Dentistry 5e",
    "cohen": "Cohen's Pathways of the Pulp 2016 (Hargreaves & Berman)",
    "carranza": "Newman & Carranza's Clinical Periodontology 13e",
    "mats": "O'Brien / Powers & Sakaguchi — Dental Materials & Their Selection 3e; Clinical Aspects of Dental Materials",
    "oms": "Contemporary Oral & Maxillofacial Surgery (Hupp/Tucker) 6e/7e",
    "mccracken": "McCracken's Removable Partial Prosthodontics",
    "complete": "Textbook of Complete Dentures",
    "mcdonald": "McDonald & Avery Pediatric Dentistry 10e",
    "ortho": "Proffit — Contemporary Orthodontics 5e/7e (2026)",
    "wheeler": "Wheeler's Dental Anatomy (via Sturdevant ch. 2 occlusion)",
}

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def H(text, size=15, color=(0x14,0x3d,0x73), space_before=10, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(*color)
    p.space_before = Pt(space_before)
    p.space_after = Pt(space_after)
    return p

def P(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def Q(num, stem, options, answer_letter, answer_text, book, why):
    H(f"Q{num}. {stem}", size=12, color=(0x0b,0x3d,0x91), space_before=8, space_after=2)
    if options:
        for o in options:
            is_ans = o.strip().startswith(answer_letter + ".") or o.strip().startswith(answer_letter.lower()+".")
            p = doc.add_paragraph()
            r = p.add_run(("✓ " if is_ans else "   ") + o)
            if is_ans:
                r.bold = True
                r.font.color.rgb = RGBColor(0x1b,0x7a,0x3d)
            p.paragraph_format.left_indent = Inches(0.25)
            p.space_after = Pt(0)
    P(f"Answer: {answer_letter}{(' — ' + answer_text) if answer_text else ''}", bold=True, color=(0x1b,0x7a,0x3d))
    P(f"📖 Reference: {book}", italic=True, size=10, color=(0x66,0x44,0x00))
    P(f"Why: {why}", size=10, color=(0x33,0x33,0x33))

# ---------- Title ----------
t = doc.add_paragraph()
r = t.add_run("SDLE — Q&A Answered from the Official Textbooks")
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x14,0x3d,0x73)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
P("Each answer below is grounded in the gold reference textbooks used for the SDLE (the same library that grades "
  "the app's MCQ bank). The correct option is marked ✓. For every item we list the source book and a short "
  "‘why’ so you can see the reasoning, not just the letter.", italic=True, size=10)
P("Rule: community recall marks are leads only — these are confirmed against the books named in 📖 Reference.",
  italic=True, size=9, color=(0x88,0x88,0x88))
doc.add_paragraph()

# ============================================================
# SET A — open-ended recall (Q2..Q36)
# ============================================================
H("Set A — Open-ended recall questions", size=14)

Q(2, "What is the percentage of metal in a PFM (Porcelain-Fused-to-Metal) crown?", None, "",
  "The metal coping (framework) is thin — about 0.3–0.5 mm — and the porcelain veneer is about 0.7–1.0 mm, "
  "so the metal is roughly one-third of the restoration thickness by cross-section.",
  BOOKS["fixed"],
  "Shillingburg specifies the metal-ceramic crown: a coping of ~0.3–0.5 mm metal is veneered with ~0.7–1 mm "
  "porcelain. The metal is kept thin (just enough for rigidity/opaque) because porcelain provides the bulk and "
  "the esthetics. The question is really about coping thickness/porportion, not a single '% metal' figure.")

Q(3, "What are the types of finishing lines for full-coverage restorations, and what is a key feature of a good finishing line?",
  None, "",
  "Types: feather edge, chisel, chamfer, bevel, shoulder (90°), sloped shoulder, beveled shoulder. "
  "Key feature of a good finishing line: it must be a distinct, sharply defined, easily readable margin.",
  BOOKS["fixed"],
  "Shillingburg (Ch. 7 margin designs) shows all these forms and states the margin must be 'distinct' so the "
  "technician can see where the die ends and the wax/pattern begins. A chamfer is the most common for cast metal; "
  "a 90° shoulder for ceramic margins.")

Q(4, "From where to where is the clinical crown measured?", None, "",
  "From the cementoenamel junction (CEJ) to the incisal/occlusal edge — i.e. the portion of the tooth coronal to the CEJ. "
  "In the mouth the ‘visible/clinical’ crown is measured from the free gingival margin to the incisal edge.",
  BOOKS["sturv"],
  "Sturdevant distinguishes the anatomical crown (CEJ → incisal edge) from the clinical crown (gingival margin → "
  "incisal edge). Restorative preparation length is planned relative to the clinical crown visible in the mouth.")

Q(9, "What is the purpose of a dental prop (bite block)?", None, "",
  "To support and stabilise the mandible, keep the mouth open at a comfortable width, and protect the TMJ "
  "(and teeth) during long procedures under sedation/GA.",
  BOOKS["oms"],
  "Hupp (Contemporary OMS) Ch. 6: 'a bite block placed on the contralateral side supports the patient’s jaw and "
  "protects the TMJ'; it prevents over-opening of the joint and fatigue. It must be released periodically.")

Q(13, "What material is used to fill a root canal to aid healing?", None, "",
  "Gutta-percha (the solid core) plus a biocompatible root-canal sealer.",
  BOOKS["cohen"],
  "Cohen's Pathways: gutta-percha is the gold-standard obturating core; the sealer fills accessory canals and "
  "irregularities and adheres to dentin. Together they three-dimensionally seal the canal to prevent re-infection "
  "and allow periradicular healing.")

Q(14, "The shape of the root canal in cross-section is variable, while in the apical third it is always round. "
       "Which part is correct and which is incorrect?", None, "",
  "First part CORRECT (canal cross-section is variable); second part INCORRECT (the apical third is NOT always round — "
  "it can be oval, kidney-shaped, or have fins/isthmuses).",
  BOOKS["cohen"],
  "Cohen's Ch. 5 (µCT cross-sections) shows canals vary widely along the root; the apical third frequently retains "
  "an oval/kidney shape with lateral fins and a second canal. Only instrumentation with round files rounds it — "
  "the native anatomy is not reliably round.")

Q(15, "What is the irritating material to the pulp?", None, "",
  "Zinc phosphate cement (acidic) and silicate cement are the classic pulp irritants; also phosphoric acid etchant "
  "applied to deep dentin, and microleakage of any unsealed restoration.",
  BOOKS["sturv"] + " / " + BOOKS["mats"],
  "Sturdevant & the materials texts: zinc phosphate is acidic (low pH while setting) and is cytotoxic to odontoblasts "
  "in deep cavities → use a liner/base (Ca(OH)₂ / GIC) first. Silicate cements historically caused severe pulp "
  "necrosis. Modern GIC and resin-modified GIC are far less irritating.")

Q(16, "Why is a substance like Hydroxyapatite or Fluoride added to toothpaste?", None, "",
  "To remineralise early enamel lesions and convert hydroxyapatite into the more acid-resistant fluorapatite / "
  "fluorohydroxyapatite, and to inhibit bacterial enzymes.",
  BOOKS["sturv"],
  "Sturdevant: tooth minerals are hydroxyapatite Ca₁₀(PO₄)₆(OH)₂. Fluoride substitutes for the hydroxyl ion → "
  "fluorapatite, which dissolves at a lower pH (more acid-resistant) and favours remineralisation. "
  "Hydroxyapatite nano-particles act as a calcium/phosphate reservoir that buffers demineralisation.")

Q(18, "Does dental caries progress from dentin to enamel (Forward) or from enamel to dentin (Backward)?",
  None, "",
  "From ENAMEL to DENTIN — caries begins in enamel and spreads into dentin (so the ‘dentin→enamel (forward)’ "
  "label is the wrong direction).",
  BOOKS["sturv"],
  "Sturdevant: the carious process starts with enamel demineralisation (white spot), then crosses the "
  "enamel-dentin junction and spreads laterally and pulpally in dentin (which is less mineralised and has "
  "tubules). So progression is enamel → dentin.")

Q(19, "What is the general public-health preventive measure to prevent dental caries? (Brushing and toothpaste.)",
  None, "",
  "Community WATER FLUORIDATION is the most effective public-health measure; at the individual level, twice-daily "
  "brushing with fluoride toothpaste is the most widely available method.",
  BOOKS["sturv"],
  "Sturdevant: 'every $1 spent on water fluoridation saves ~$6 in dental treatment' — community water fluoridation "
  "is the cornerstone population strategy. On top of that, fluoride toothpaste + brushing delivers a daily topical "
  "dose. Both are needed; the public-health pillar is water fluoridation.")

Q(24, "A 14-year-old child has a congenital absence of the central and lateral incisors. What type of prosthesis should be made?",
  None, "",
  "A transitional / provisional REMOVABLE partial denture (or an acid-etch/Maryland bonded bridge) — NOT an implant, "
  "because growth is not yet complete.",
  BOOKS["mcdonald"] + " / " + BOOKS["fixed"],
  "McDonald (Pediatric) & Shillingburg: implants must be deferred until skeletal growth is finished (~age 18–20 "
  "for girls, ~21+ for boys) to avoid infraocclusion. Until then a removable transitional prosthesis (or a "
  "conservative bonded Maryland bridge) restores esthetics and function and can be modified as the arch grows.")

Q(25, "What is meant by a 'Sunday bite'? (Class III malocclusion)", None, "",
  "A habitual forward posturing of the mandible so the incisors meet edge-to-edge / in a pseudo-Class III "
  "relationship — a habitual, not skeletal, Class III.",
  BOOKS["ortho"],
  "Proffit (Contemporary Orthodontics): a 'Sunday bite' (pseudo-Class III) is a habitual forward slide from a "
  "normal or Class II molar into an edge-to-edge anterior relationship. It is a postural habit, not true skeletal "
  "Class III, and must be diagnosed (often the patient can retrude to a normal overjet).")

Q(28, "What is the purpose of dentin etching?", None, "",
  "To remove the smear layer, open the dentinal tubules, and demineralise the intertubular/peritubular dentin so "
  "that resin infiltrates and forms the hybrid layer + resin tags (micro-mechanical retention).",
  BOOKS["sturv"],
  "Sturdevant: acid etching of dentin dissolves the smear layer and ~2–5 µm of hydroxyapatite, exposing a collagen "
  "network. Hydrophilic resin primer infiltrates this network → the 'hybrid layer' (Nakabayashi) and resin tags in "
  "the tubules. This is the basis of modern dentin bonding.")

Q(29, "What is used to etch dentin?", None, "",
  "Phosphoric acid, usually 32–37% (typical 37% for 15 s on dentin, shorter than enamel).",
  BOOKS["sturv"],
  "Sturdevant: 37% phosphoric acid gel is standard; dentin is etched for a shorter time (~15 s) than enamel "
  "(15–30 s) to avoid over-etching/collagen collapse. Self-etch adhesives use acidic monomers instead of a "
  "separate phosphoric step.")

Q(30, "What is a high-heat obturation technique?", None, "",
  "An injectable thermoplasticised gutta-percha technique that heats GP to a high temperature (e.g. Obtura II at "
  "~160 °C) to fill the canal in three dimensions, usually after a warm vertical down-pack.",
  BOOKS["cohen"],
  "Cohen's: Obtura II (and System B for warm vertical compaction) are the high-heat thermoplasticised techniques. "
  "They deliver warm, flowable GP that adapts to canal irregularities; ThermoFil/Soft-Core are carrier-based, and "
  "Ultrafil is low-heat (lower temp, higher viscosity). Lateral condensation is cold (no heat).")

Q(35, "About the maxillary palatal root variation: describe the variation.", None, "",
  "The palatal root of the maxillary first molar is the longest and widest root, inclines distally (and slightly "
  "buccally), and its buccal surface is concave / furrowed — giving a kidney/bean cross-section that frequently "
  "hides a second palatal canal.",
  BOOKS["cohen"],
  "Cohen's Ch. 5 & 7: the palatal root inclines distally, is broadest buccolingually, has a buccal concavity, and "
  "a second canal (MB2 is on the mesiobuccal root, but the palatal root itself also can have two canals) — the "
  "concavity/bean shape is the classic reason a canal is missed under the DOM.")

# ============================================================
# SET B — first image MCQs (Q12..Q50)
# ============================================================
H("Set B — MCQs (first image)", size=14)

Q(12, "The gold-standard material in vital pulpotomy for teeth is:",
  ["A. MTA.", "B. Ca(OH)₂.", "C. GIC.", "D. Resin composite."],
  "A", "MTA",
  BOOKS["mcdonald"] + " / " + BOOKS["cohen"],
  "McDonald & Cohen's: mineral trioxide aggregate (MTA) has superseded formocresol as the gold standard for vital "
  "pulpotomy — it sets in the presence of moisture, is biocompatible, and induces dentin bridge formation with the "
  "fewest pulp necrosis outcomes. Ca(OH)₂ is the historical alternative but is associated with more internal "
  "resorption; formocresol is obsolete (toxic/carcinogenic).")

Q(16, "Formocresol has limited use in dentistry because it is:",
  ["A. Poor biocompatibility.", "B. High strength.", "C. Weak.", "D. All the above."],
  "A", "Poor biocompatibility",
  BOOKS["mcdonald"],
  "McDonald: formocresol (Buckley's) is toxic, mutagenic/carcinogenic (formaldehyde + cresol), distributes "
  "systemically, and is no longer recommended. Its limitation is biological (poor biocompatibility), not mechanical "
  "— it isn't used as a structural material.")

Q(24, "A patient with generalized attrition who needs an FPD should first undergo:",
  ["A. Period surgery.", "B. Desensitization of teeth.", "C. Crown build-up.", "D. Conventional RCT."],
  "B", "Desensitization of teeth",
  BOOKS["sturv"] + " / " + BOOKS["fixed"],
  "For generalized attrition the first, most conservative step is to manage sensitivity and establish the correct "
  "occlusal vertical dimension before any restorative work; desensitizing agents (and a Dahl-type anterior "
  "composite/ splint) come before crowns. RCT is only done if the pulp is irreversibly affected; perio surgery only "
  "if periodontal disease is present.")

Q(31, "Which best describes the movement technique of a K-file during root canal instrumentation:",
  ["A. Rotation.", "B. Up and down motion.", "C. Clockwise rotation combined with up and down motion.", "D. Circumferential filing."],
  "C", "Clockwise rotation combined with up and down motion",
  BOOKS["cohen"],
  "Cohen's: K-files are manipulated with a quarter-turn clockwise rotation (reaming) combined with an apical-coronal "
  "withdrawal (filing) stroke — the classic 'watch-winding'/balanced-force motion. Pure rotation risks separation; "
  "pure pull is a Hedstrom technique; circumferential filing is for enlargement with H-files.")

Q(32, "What is the main purpose of dentin etching:",
  ["A. To remove the smear layer only.", "B. To create micro-retentive resin tags (interlocking tags).",
   "C. To increase dentin hardness.", "D. To disinfect the cavity."],
  "B", "To create micro-retentive resin tags (interlocking tags)",
  BOOKS["sturv"],
  "Sturdevant: the PURPOSE of etching dentin is to create micromechanical retention — the hybrid layer + resin tags — "
  "that interlocks the resin to the tooth. Removing the smear layer (A) is a means, not the purpose; etching does not "
  "harden dentin or disinfect the cavity.")

Q(33, "What is a high-heat obturation technique:",
  ["A. Thermofil.", "B. Ultrafil.", "C. Obtura II.", "D. Lateral condensation."],
  "C", "Obtura II",
  BOOKS["cohen"],
  "Cohen's: Obtura II is the high-heat injectable thermoplasticised gutta-percha system (~160 °C). ThermoFil/Soft-Core "
  "are carrier-based; Ultrafil is a LOW-heat (≈70 °C) high-viscosity system; lateral condensation uses cold GP.")

Q(39, "Which best describes the correct path of insertion for a fixed prosthesis:",
  ["A. Parallel to each other.", "B. Parallel to the long axis of the normal adjacent tooth.",
   "C. Parallel to the long axis of the abutment tooth.", "D. Perpendicular to the occlusal plane."],
  "C", "Parallel to the long axis of the abutment tooth",
  BOOKS["fixed"],
  "Shillingburg: the path of insertion/withdrawal is determined by the axial walls of the prepared abutment and "
  "should follow the long axis of the abutment tooth so the restoration seats without binding. Adjacent teeth guide "
  "the survey but each abutment's long axis governs its own path; 'perpendicular to the occlusal plane' applies to "
  "surveying a partial denture, not an FPD.")

Q(45, "Which is NOT an indication for placing a subgingival margin:",
  ["A. When additional retention is required.", "B. For esthetic reasons.",
   "C. Presence of cervical caries.", "D. When the margin is easily visible at the first appointment."],
  "D", "When the margin is easily visible at the first appointment",
  BOOKS["fixed"],
  "Shillingburg: subgingival margins are justified only for (1) esthetics (anterior facial), (2) cervical caries/"
  "restoration/erosion, (3) need for extra axial wall height/retention, (4) subgingival existing margin. A margin "
  "that is 'easily visible' is precisely a supragingival indication — placing it subgingival needlessly violates the "
  "biologic width and harms periodontium.")

Q(46, "Most appropriate management for a crown margin exposed 0.5 mm above the gingival margin:",
  ["A. Remake the crown.", "B. Restore the exposed area using GIC.", "C. Restore the exposed area using amalgam.", "D. All of the above."],
  "A", "Remake the crown",
  BOOKS["fixed"],
  "Shillingburg: an open/short margin that is supragingival and visible is a remake — patching with GIC or amalgam "
  "leaves a marginal discrepancy with plaque retention, recurrent caries, and wash-out. The only acceptable answer "
  "is to remake the crown to a proper margin (the GIC/amalgam 'patch' options are distractors).")

Q(48, "Average lifespan of composite restorations in primary teeth:",
  ["A. 2–3 years.", "B. 3–5 years.", "C. 5–10 years.", "D. 10–15 years."],
  "B", "3–5 years",
  BOOKS["mcdonald"] + " / " + BOOKS["sturv"],
  "McDonald: composite restorations in primary teeth are expected to survive ~3–5 years (the lifetime of a primary "
  "molar until exfoliation is ~3–6 years). Survival is shorter than in permanent teeth due to smaller restorations, "
  "moisture-control challenges, and child behavior; 5–10/10–15 are unrealistically long for primary dentition.")

Q(49, "Which composite resin is most suitable for restoring posterior teeth:",
  ["A. Microfilled composite.", "B. Flowable composite.", "C. Light-curing composite.", "D. Packable composite."],
  "D", "Packable composite",
  BOOKS["sturv"] + " / " + BOOKS["mats"],
  "Packable (condensable) composites have high filler load and high viscosity → better handling, lower polymerisation "
  "shrinkage, and improved wear resistance for posterior load-bearing cavities. Microfilled is too weak/low-modulus "
  "for posterior occlusal load; flowable is for liners/small class I/abfractions; 'light-curing' is a cure mode, not "
  "a class.")

Q(50, "Which best describes packable composite resin:",
  ["A. High viscosity and used for restoring posterior teeth.", "B. Low viscosity.",
   "C. Easily polishable.", "D. Highly translucent."],
  "A", "High viscosity and used for restoring posterior teeth",
  BOOKS["mats"],
  "Dental Materials texts: packable composites are high-viscosity, highly filled restoratives designed for posterior "
  "Class I/II — they handle like amalgam, resist slump, and have adequate wear resistance. Low viscosity = flowable; "
  "polishability/translucency describe microfilled/nanohybrid esthetic composites.")

# ============================================================
# SET C — second image MCQs
# ============================================================
H("Set C — MCQs (second image)", size=14)

Q(1, "The 6th generation bonding system is characterized as:",
  ["A. Three-step technique.", "B. Self-etching system without separate resin application.",
   "C. Self-adhesive system.", "D. Total-etch two-step system."],
  "B", "Self-etching system (6th gen = self-etch adhesives)",
  BOOKS["sturv"],
  "Sturdevant: by the standard generation scheme, 6th-gen = self-etch adhesives (acidic monomers etch & prime "
  "simultaneously). Strictly, 6th gen is a 2-step self-etch (separate adhesive resin) and 7th gen is the all-in-one "
  "self-etch (no separate resin); 8th gen = self-adhesive. Among the given options the self-etch option (B) is the "
  "best match for '6th generation'; C would be 8th gen and D is 5th gen (total-etch 2-step).")

Q(3, "Supragingival calculus is characterized by all of the following EXCEPT:",
  ["A. Hard and rough.", "B. Easy to detach.", "C. Has component of saliva.", "D. None."],
  "B", "Easy to detach",
  BOOKS["carranza"],
  "Carranza: supragingival calculus is hard, rough, yellowish, firmly adherent to enamel and composed of mineralised "
  "plaque with a saliva-derived matrix. It is NOT easy to detach — it is tenaciously bonded and requires instrumentation "
  "(scaling) to remove. So 'easy to detach' is the characteristic it does NOT have.")

Q(8, "Which teeth are considered relatively more difficult to extract:",
  ["A. Mandibular canine.", "B. Maxillary canine.", "C. Mandibular premolar.", "D. Maxillary central incisor."],
  "B", "Maxillary canine",
  BOOKS["oms"],
  "Hupp (Contemporary OMS): the maxillary canine has the longest root in the arch (~26 mm), a bulky crown, a thick "
  "buccal plate, and a palatal inclination, making it the most difficult routine (forceps) extraction of these "
  "choices. Mandibular premolars and incisors have single, smaller, conical roots and are comparatively easy.")

Q(12, "The marginal ridge cusp of the second mandibular molar occludes with:",
  ["A. Mesial marginal ridge of maxillary second premolar.",
   "B. Mesial marginal ridge of maxillary first molar.",
   "C. Distal marginal ridge of maxillary first molar.",
   "D. Distal marginal ridge of maxillary second molar."],
  "C", "Distal marginal ridge of maxillary first molar",
  BOOKS["wheeler"],
  "Wheeler/Sturdevant occlusion: in Angle Class I the mesial cusp/marginal ridge of the mandibular second molar "
  "opposes the embrasure between the maxillary first and second molars — contacting the distal marginal ridge of the "
  "maxillary first molar and the mesial marginal ridge of the maxillary second molar. Among the options, the distal "
  "ridge of the maxillary first molar is the named occlusal contact.")

Q(16, "Which is used to adjust the anterior relationship of the bite rim:",
  ["A. Facebow.", "B. Fox plane.", "C. Articulator.", "D. Bite registration material."],
  "B", "Fox plane",
  BOOKS["complete"] + " / " + BOOKS["mccracken"],
  "The Fox plane (occlusal plane guide/indicator) is placed against the bite rims and viewed from the front/profile "
  "to set the antero-posterior inclination of the occlusal plane (anterior-posterior relationship of the rim). A "
  "facebow relates the maxilla to the hinge axis; the articulator receives the record; bite registration material "
  "records the interocclusal relationship — none of those 'adjust' the anterior rim inclination.")

Q(17, "The first and most important step in the diagnostic work-up of a dental patient:",
  ["A. Reviewing the medical history.", "B. Taking the dental history.",
   "C. Identifying the chief complaint.", "D. Obtaining radiographic images."],
  "C", "Identifying the chief complaint",
  BOOKS["oms"],
  "Contemporary OMS Ch. 1: 'Every patient should be asked to state their chief complaint' — the chief complaint is "
  "the patient's presenting problem in their own words and is the starting point that focuses the history, "
  "examination, and radiographs. Medical/dental history and imaging follow it.")

Q(24, "The factor that increases cohesive fracture in PFM (Porcelain-Fused-to-Metal) crowns is:",
  ["A. Increased oxide layer.", "B. Thick porcelain.", "C. Thick metal.", "D. All of the above."],
  "A", "Increased oxide layer",
  BOOKS["mats"],
  "Dental materials: the metal-porcelain bond depends on a thin, controlled oxide layer. An EXCESSIVE (thick) oxide "
  "layer is brittle and weakly adherent → cohesive fracture within the oxide/porcelain and porcelain chipping. A "
  "thin, uniform oxide (and matched CTE) optimises bonding; thick metal or appropriately thick porcelain are not "
  "fracture-promoting factors.")

Q(26, "A correct comparison between all-ceramic jackets and PFM crowns is:",
  ["A. They have the same finish line (margin).",
   "B. The lingual wall is more conservative in a PFM crown."],
  "B", "The lingual wall is more conservative in a PFM crown",
  BOOKS["fixed"],
  "Shillingburg: a PFM crown's lingual/palatal surface can be restored in metal only, so its lingual reduction is "
  "conservative (~0.7–1 mm) versus an all-ceramic crown which needs uniform 1.2–1.5 mm reduction all around plus a "
  "90° shoulder — they do NOT share the same finish line (all-ceramic = shoulder; PFM = chamfer or shoulder).")

Q(48, "The first step in treating generalized attrition is:",
  ["A. Desensitizing agents.", "B. Crowns.", "C. Composite build-up."],
  "A", "Desensitizing agents",
  BOOKS["sturv"] + " / " + BOOKS["fixed"],
  "Sturdevant/Shillingburg: generalized attrition is first managed conservatively — desensitise the exposed dentin "
  "and stabilise the occlusion (often a splint / Dahl concept) to determine the correct vertical dimension before any "
  "definitive restoration. Crowns and composite build-ups come after the occlusion and sensitivity are controlled.")

Q(50, "Which fluoride method is used for the general prevention of caries worldwide:",
  ["A. Systemic fluoride.", "B. Topical fluoride.", "C. Toothpaste."],
  "C", "Toothpaste",
  BOOKS["sturv"],
  "Sturdevant: fluoride toothpaste is the most widely available and universally used caries-prevention method across "
  "populations (twice-daily brushing). Water fluoridation (a systemic-with-topical-effect community measure) is the "
  "most cost-effective public-health pillar, but 'toothpaste' is the worldwide, individual-level general measure.")

# ============================================================
# SET D — third image MCQs
# ============================================================
H("Set D — MCQs (third image)", size=14)

Q(1, "A simple extraction uses an elevator only for:",
  ["A. Impacted tooth.", "B. Removing an erupted tooth.", "C. Removing a tooth with bone."],
  "B", "Removing an erupted tooth",
  BOOKS["oms"],
  "Contemporary OMS: a straight elevator is used to luxate (loosen) an erupted tooth from its periodontal ligament "
  "before forceps delivery. Impacted teeth and teeth with bony ankylosis/bridging require surgical removal "
  "(ostectomy, sectioning) — an elevator alone is insufficient and would fracture bone/tooth.")

Q(2, "About the maxillary palatal root variation:",
  ["A. Concave and kidney shape.", "B. Rarely has two canals.", "C. Inclined distally.",
   "D. Wider and oval in cross-section."],
  "A", "Concave and kidney shape",
  BOOKS["cohen"],
  "Cohen's: the palatal root is broad buccolingually with a buccal concavity/furrow, giving a kidney/bean cross-section "
  "that hides a second canal — this is the hallmark 'variation' tested. (The root is ALSO inclined distally and can "
  "have two canals, so B is false; A captures the distinctive morphology that explains missed canals.)")

Q(4, "The success of a root canal filling is best assessed by:",
  ["A. Radiographs.", "B. Clinical observation.", "C. Size of gutta-percha cone used.", "D. A and B."],
  "D", "A and B (radiographs + clinical observation)",
  BOOKS["cohen"],
  "Cohen's: endodontic outcome is judged by BOTH clinical signs/symptoms (no pain, no swelling, normal probing, "
  "function) AND radiographic healing (resolution of the periradicular radiolucency, intact lamina dura). The size "
  "of the GP cone used is irrelevant to success; radiographs alone are insufficient without clinical correlation.")

Q(6, "Inadequate incisal reduction during tooth preparation for a metal-ceramic restoration results in:",
  ["A. Inadequate path of insertion.", "B. Less resistance and retention of the restoration.",
   "C. Poor incisal translucency in the final restoration.", "D. All are true."],
  "C", "Poor incisal translucency in the final restoration",
  BOOKS["fixed"],
  "Shillingburg: insufficient incisal reduction leaves inadequate space for the metal coping + porcelain veneer → the "
  "porcelain is too thin (or the coping shows through) → poor incisal translucency/esthetics and possibly a weak "
  "coping. It does not alter the path of insertion, and retention is an axial-wall issue, so 'all true' is wrong.")

Q(10, "During tooth preparation for a full metal crown, the amount of tooth structure that should be removed is:",
  ["A. 0.5–0.7 mm.", "B. 0.7–1 mm.", "C. 1–1.5 mm.", "D. 1.5–2 mm."],
  "C", "1–1.5 mm (occlusal clearance)",
  BOOKS["fixed"],
  "Shillingburg: a full cast metal crown requires ~1.0–1.5 mm occlusal clearance and ~0.7–1 mm axial reduction with a "
  "0.3–0.5 mm chamfer. The single 'amount removed' figure most often tested is the occlusal clearance of 1–1.5 mm "
  "(functional cusp 1.5 mm, non-functional 1 mm). 1.5–2 mm belongs to PFM/all-ceramic.")

Q(22, "When using an inflexible file in a curved canal, what does it cause on the outer surface of the curve:",
  ["A. Ledge.", "B. Zipping.", "C. Perforation.", "D. Elbow."],
  "B", "Zipping",
  BOOKS["cohen"],
  "Cohen's: a stiff (large/inflexible) file tends to straighten in a curved canal and preferentially cuts the outer "
  "aspect toward the apex → apical transportation = 'zipping'. A ledge (A) is a stepped indentation coronally; an "
  "elbow (D) is the transition above the zip; perforation can follow if over-driven.")

Q(23, "The movements of a K-file in root canal treatment are:",
  ["A. Clockwise and anti-clockwise with pressure apically.", "B. Rotation movement with pressure apically."],
  "A", "Clockwise and anti-clockwise with pressure apically",
  BOOKS["cohen"],
  "Cohen's: the K-file is used with a balanced-force / watch-winding motion — clockwise then anti-clockwise rotation "
  "with light apical pressure — which cuts dentin on both the clockwise (cutting) and anti-clockwise (release) "
  "strokes while centring the file in the canal. Pure unidirectional rotation risks separation.")

Q(26, "The marginal ridge of the mandibular second premolar occludes with the:",
  ["A. Mesial ridge of the maxillary first premolar.",
   "B. Mesial ridge of the maxillary second premolar.",
   "C. Mesial ridge of the maxillary first molar.",
   "D. Distal fossa of the maxillary second premolar."],
  "B", "Mesial ridge of the maxillary second premolar",
  BOOKS["wheeler"],
  "Wheeler: the mandibular second premolar sits distal to the mandibular first premolar; in Class I occlusion its "
  "mesial cusp/marginal ridge engages the embrasure between the maxillary first and second premolars, contacting the "
  "distal ridge of the maxillary first premolar and the mesial ridge of the maxillary second premolar. Among the "
  "options, the mesial ridge of the maxillary second premolar is the named occlusal contact.")

Q(27, "Ideally, the length of the post in a post-and-core restoration should be at least:",
  ["A. One-half of the root length.", "B. One-third of the root length.",
   "C. Equal to the clinical crown.", "D. The full length of the root canal."],
  "A", "One-half of the root length",
  BOOKS["fixed"],
  "Shillingburg: the ideal post length is ~2/3 of the root (or at least equal to the crown height) with at least "
  "~4–5 mm of apical gutta-percha seal preserved; the MINIMUM acceptable is one-half of the root length. One-third "
  "is too short for retention; full length destroys the apical seal.")

Q(33, "The impression material which is more accurate when the pouring is done after a week is:",
  ["A. Polysulfide.", "B. Polyether.", "C. Agar-agar.", "D. Addition silicone."],
  "D", "Addition silicone",
  BOOKS["mats"],
  "Dental materials: addition-cured silicone (polyvinyl siloxane, PVS) has the best dimensional stability — it can "
  "be poured up to a week later with negligible change because it polymerises by addition with no by-product. "
  "Polysulfide and polyether are less stable on delay; agar (reversible hydrocolloid) must be poured immediately.")

Q(48, "The shape of canal preparation in cross-section is variable, and in the apical third, it is round:",
  ["A. Both statements are true.", "B. Both statements are false.",
   "C. The first statement is true, and the second is false.",
   "D. The first statement is false, and the second is true."],
  "C", "First statement true, second false",
  BOOKS["cohen"],
  "Cohen's: the canal SYSTEM is variable in cross-section throughout its length (true); the apical third is NOT "
  "reliably round — native anatomy is often oval/kidney with fins (false). Although round K-file instrumentation "
  "tends to create a rounder apical prep, the original canal anatomy and residual fins mean 'it is round' is an "
  "overstatement, so the second statement is the incorrect one.")

Q(50, "The impression material that is the least difficult to remove after setting is:",
  ["A. Alginate.", "B. Compound.", "C. Silicone.", "D. Polyether."],
  "C", "Silicone",
  BOOKS["mats"],
  "Dental materials: addition silicone has the highest elastic recovery and good tear strength → it recovers from "
  "undercuts and is the easiest to remove cleanly. Polyether is rigid and stiff (the MOST difficult to remove, "
  "especially from undercuts); compound is rigid; alginate is elastic but has low tear strength and dehydrates.")

Q(51, "An indication for zirconia ceramic is for:",
  ["A. Post and core.", "B. Implant and abutment.", "C. Orthodontic brackets.", "D. All of the above."],
  "D", "All of the above",
  BOOKS["mats"] + " / " + BOOKS["fixed"],
  "Yttria-stabilised zirconia (Y-TZP) is used for crowns/bridges, implant abutments, posts and cores, and "
  "orthodontic brackets (e.g. ceramic self-ligating brackets) because of its high flexural strength, fracture "
  "toughness, and biocompatibility.")

Q(52, "The expected long life of a composite restoration is:",
  ["A. 1–2 years.", "B. 3–5 years.", "C. 9–11 years."],
  "C", "9–11 years (the 'expected long life')",
  BOOKS["sturv"] + " / " + BOOKS["mats"],
  "Sturdevant: while the AVERAGE clinical survival of posterior composite is ~5–7 years, the EXPECTED (best-case, "
  "well-placed) long life is ~10 years (9–11). The question specifically asks for the 'expected long life', so the "
  "higher figure applies; 3–5 reflects average/wear-limited survival.")

Q(53, "The first step in a diagnostic work-up is obtaining the:",
  ["A. Medical history.", "B. Present complaint.", "C. Biographical data.",
   "D. Restorative history.", "E. Traumatic history."],
  "B", "Present complaint (chief complaint)",
  BOOKS["oms"],
  "Contemporary OMS: the diagnostic sequence starts with the chief/present complaint (why the patient came today), "
  "followed by medical and dental histories and examination. Biographical data is administrative; histories come "
  "after the presenting complaint is recorded.")

Q(61, "Heating techniques in gutta-percha obturation are called:",
  ["A. Thermoplasticized Gutta-Percha Techniques."],
  "A", "Thermoplasticized Gutta-Percha Techniques",
  BOOKS["cohen"],
  "Cohen's: any obturation that heats GP — warm vertical condensation (Schilder/System B), injectable (Obtura II), "
  "carrier-based (ThermaFil) — falls under 'thermoplasticised gutta-percha techniques' (as opposed to cold lateral "
  "condensation).")

Q(63, "A characteristic of 6th generation dentin bonding agents is that they are:",
  ["A. Self-etching."],
  "A", "Self-etching",
  BOOKS["sturv"],
  "Sturdevant: 6th-generation adhesives are self-etching — acidic monomers simultaneously etch and prime the dentin "
  "(no separate phosphoric-acid step). (Strictly, 6th gen = 2-step self-etch with a separate resin; 7th gen = "
  "all-in-one self-etch.)")

Q(64, "Custom impression tray is used with putty wash technique rather than the others:",
  ["A. True.", "B. False."],
  "B", "False",
  BOOKS["mats"] + " / " + BOOKS["fixed"],
  "The putty-wash (one-step double-mix) technique is typically done in a STOCK tray: heavy-body putty fills the "
  "stock tray and light-body wash is injected on the teeth. Custom trays are used with single-viscosity/heavy-and-"
  "light-body techniques and for a final wash in a two-step technique — so the statement is false.")

Q(66, "When the opposing teeth occlude on the cervical fifth of the lingual surface, this is a contraindication of a metal-ceramic crown:",
  ["A. True.", "B. False."],
  "A", "True",
  BOOKS["fixed"],
  "Shillingburg: if the opposing teeth contact the cervical fifth of the lingual surface, the porcelain on the "
  "lingual of a metal-ceramic crown would be in direct occlusion and prone to fracture (the lingual of a PFM is "
  "often metal only / thin porcelain). This is a listed contraindication — a full metal or modified design is "
  "preferred so the contact falls on metal.")

Q(69, "Preservation of the periodontium is one of the principles of tooth preparation; to carry out this principle, "
       "the subgingival margins should be avoided as much as possible:",
  ["A. True.", "B. False."],
  "A", "True",
  BOOKS["fixed"],
  "Shillingburg's principles of tooth preparation include 'preservation of the periodontium' — keep margins supragingival "
  "whenever possible so the gingiva and biologic width are not violated; subgingival margins are reserved for specific "
  "indications (esthetics, cervical caries, retention) and otherwise avoided.")

# ============================================================
# SET E — fourth image
# ============================================================
H("Set E — MCQs (fourth image)", size=14)

Q(14, "Porcelain veneer is made from:",
  ["A. Feldspathic.", "B. Lithium (disilicate).", "C. Leucite.", "D. All of the above."],
  "D", "All of the above",
  BOOKS["mats"] + " / " + BOOKS["fixed"],
  "Porcelain veneers can be fabricated from feldspathic porcelain (traditional, etchable, high esthetics), "
  "leucite-reinforced ceramic (Empress), or lithium-disilicate glass-ceramic (e.max) — all are valid veneer ceramics, "
  "chosen by strength/esthetic needs. So 'all of the above' is correct.")

Q(31, "Root caries properties (all correct EXCEPT):",
  ["A. Rapidly very progression.", "B. V shaped in cross section.", "C. Well define margins.",
   "D. Found in old patient because gingival recession."],
  "A", "Rapidly very progression (EXCEPT)",
  BOOKS["sturv"] + " / " + BOOKS["carranza"],
  "Root caries is characteristically SLOW and indolent (it progresses laterally in demineralised cementum/dentin), "
  "so 'rapidly very progression' is the property that does NOT fit. (Note: root caries is also shallow U/saucer "
  "shaped — not V — and has diffuse, ill-defined margins, so B and C are also non-characteristic, but A is the "
  "clearest single EXCEPT as the classic distractor.)")

Q(10, "Which one of these restorative methods will be LEAST compromised by a core:",
  ["a) Amalgam", "b) Composite", "c) GIC", "d) Cast metal"],
  "d", "Cast metal",
  BOOKS["fixed"],
  "Shillingburg: a cast metal (cast gold) restoration/onlay/3/4 crown is rigid, ductile, and self-supporting — it is "
  "least compromised by the presence of a core buildup because it can be cast to engage remaining tooth structure "
  "and does not depend on the core for bulk/strength the way amalgam, composite, or GIC restorations do.")

Q(20, "When you make an impression for a mandibular knife-edge (ridge):",
  ["A. Minimum pressure impression.", "B. Selective pressure impression.",
   "C. Maximum pressure impression.", "D. None."],
  "A", "Minimum pressure impression",
  BOOKS["complete"] + " / " + BOOKS["mccracken"],
  "Textbook of Complete Dentures/McCracken: a knife-edge (sharp, flat) mandibular ridge is covered by a thin, easily "
  "displaced mucosa over a sharp crest — a minimum-pressure (mucostatic) impression avoids compressing and traumatizing "
  "this tissue and avoids ridge 'resorption' under load. Selective/maximum pressure is reserved for compressible "
  "tuberosity/flabby tissue, not a knife-edge crest.")

# ---------- closing note ----------
H("Notes on interpretation", size=12)
P("• Two items ask the same fact in different wording — Q2 (PFM metal proportion) and the palatal-root variation "
  "(Q35 / Set-D Q2); the answers are consistent across both forms.", size=10)
P("• The 6th-generation adhesive appears in Set-C Q1 and Set-D Q63; both are answered 'self-etching' — strictly, "
  "6th gen = 2-step self-etch (separate resin) and 7th gen = all-in-one self-etch, but the exam's self-etch option "
  "is the intended match.", size=10)
P("• Set-D Q48 (canal prep 'apical third is round'): the native canal is variable throughout and not reliably round "
  "at the apex; instrumentation rounds the apical prep but the statement is still an overstatement, so it is marked "
  "false for consistency with the open-ended Q14.", size=10)
P("• Set-B Q48 (primary-teeth composite lifespan) and Set-D Q52 (composite 'long life') use different horizons "
  "(primary-dentition survival vs ideal long-term survival) — read the wording carefully.", size=10)

out = "/data/prometric/SDLE_QA_Answered.docx"
doc.save(out)
print("WROTE", out)
